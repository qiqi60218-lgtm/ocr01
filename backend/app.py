"""
DocVision AI 后端服务

- Flask 提供静态页与 REST API：旋转、OCR、表格识别、文本纠错、LLM抽取、保存Word。
- 支持本地 Tesseract 与远程（豆包/网关）OCR，配置从 `.env` 读取。
"""
from flask import Flask, request, jsonify, send_from_directory, send_file
import cv2
import numpy as np
import pytesseract
from pytesseract import Output
from PIL import Image, ImageEnhance
import re
import os
import logging
import base64
import time
from io import BytesIO
from flask_cors import CORS
# 用于生成Word文档
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv

# 配置日志
logging.basicConfig(level=logging.INFO,
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger('DocVision')

# 加载环境变量
load_dotenv()

# 创建Flask应用
app = Flask(__name__, static_folder='../frontend', static_url_path='/')
CORS(app, resources={r"/*": {"origins": "*"}})  # 更宽松的跨域配置

# 配置Tesseract路径，增加自动查找功能
def find_tesseract_path():
    # 常见的Tesseract安装路径
    possible_paths = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        r'D:\Program Files\Tesseract-OCR\tesseract.exe'
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            logger.info(f"找到Tesseract: {path}")
            return path
    
    logger.warning("未找到Tesseract，将使用系统PATH中的默认路径")
    return 'tesseract'  # 希望在PATH中

pytesseract.pytesseract.tesseract_cmd = find_tesseract_path()

# 核心OCR处理类
class DocRecognizer:
    def __init__(self):
        # 不再使用tempfile，减少系统资源占用
        self.supported_formats = ['jpg', 'jpeg', 'png', 'bmp']
        logger.info("DocRecognizer初始化完成")
    
    def crop_image(self, img, crop_area):
        """根据框选区域裁剪图像"""
        if img is None:
            raise ValueError("输入图像为空")
        
        if not crop_area:
            # 如果没有裁剪区域参数，返回原始图像
            logger.info("没有提供裁剪区域参数，使用原图")
            return img
        
        try:
            h, w = img.shape[:2]
            
            # 获取裁剪坐标，不再依赖active标志
            x1 = max(0, int(crop_area.get('x1', 0)))
            y1 = max(0, int(crop_area.get('y1', 0)))
            x2 = min(w, int(crop_area.get('x2', w)))
            y2 = min(h, int(crop_area.get('y2', h)))
            
            # 不再自动判断是否为全图选择，总是按照用户提供的坐标进行裁剪
            # 这样可以确保用户的框选操作总是被尊重
            logger.info(f"准备裁剪图像，区域: ({x1}, {y1}) to ({x2}, {y2})")
                
            # 确保裁剪区域有效
            if x1 >= x2 or y1 >= y2:
                logger.warning(f"无效的裁剪区域: ({x1}, {y1}) to ({x2}, {y2})")
                return img
            
            # 裁剪图像
            cropped = img[y1:y2, x1:x2]
            logger.info(f"图像裁剪完成，裁剪区域: ({x1}, {y1}) to ({x2}, {y2})，尺寸: {cropped.shape}")
            return cropped
        except Exception as e:
            logger.error(f"图像裁剪失败: {str(e)}")
            # 如果裁剪失败，返回原始图像
            return img
    
    def recognize_table(self, image, crop_area=None):
        """识别表格并提取表格内容"""
        try:
            # 先应用裁剪区域（如果提供）
            if crop_area and isinstance(crop_area, dict):
                # 获取裁剪坐标
                x1 = int(crop_area.get('x1', 0))
                y1 = int(crop_area.get('y1', 0))
                x2 = int(crop_area.get('x2', image.shape[1]))
                y2 = int(crop_area.get('y2', image.shape[0]))
                
                # 确保坐标有效
                x1 = max(0, x1)
                y1 = max(0, y1)
                x2 = min(image.shape[1], x2)
                y2 = min(image.shape[0], y2)
                
                # 确保x2 > x1且y2 > y1
                if x2 > x1 and y2 > y1:
                    logger.info(f"应用裁剪区域: x1={x1}, y1={y1}, x2={x2}, y2={y2}")
                    image = image[y1:y2, x1:x2]
                else:
                    logger.warning("无效的裁剪区域，将使用完整图像")
            elif crop_area:
                # 向后兼容，使用现有的crop_image方法
                image = self.crop_image(image, crop_area)
            
            # 预处理图像以增强表格线条
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY) if len(image.shape) == 3 else image
            
            # 应用高斯模糊减少噪声
            blurred = cv2.GaussianBlur(gray, (3, 3), 0)
            
            # 使用自适应阈值化
            thresh = cv2.adaptiveThreshold(
                blurred, 
                255, 
                cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                cv2.THRESH_BINARY_INV, 
                11, 
                2
            )
            
            # 检测水平线
            horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
            detect_horizontal = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, horizontal_kernel, iterations=2)
            horizontal_lines = cv2.HoughLinesP(
                detect_horizontal, 
                1, 
                np.pi/180, 
                threshold=100, 
                minLineLength=100, 
                maxLineGap=10
            )
            
            # 检测垂直线
            vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
            detect_vertical = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, vertical_kernel, iterations=2)
            vertical_lines = cv2.HoughLinesP(
                detect_vertical, 
                1, 
                np.pi/180, 
                threshold=100, 
                minLineLength=100, 
                maxLineGap=10
            )
            
            # 优先尝试基于网格线的结构化表格提取
            try:
                def _cluster_positions(vals, tol=5):
                    vals = sorted(vals)
                    clusters = []
                    for v in vals:
                        if not clusters or abs(v - clusters[-1][-1]) > tol:
                            clusters.append([v])
                        else:
                            clusters[-1].append(v)
                    # 使用平均值代表每个聚类
                    return [int(sum(c) / len(c)) for c in clusters]
                # 提取水平线的y坐标与垂直线的x坐标
                h_pos_candidates = []
                if horizontal_lines is not None:
                    for line in horizontal_lines:
                        x1, y1, x2, y2 = line[0]
                        h_pos_candidates.extend([y1, y2])
                v_pos_candidates = []
                if vertical_lines is not None:
                    for line in vertical_lines:
                        x1, y1, x2, y2 = line[0]
                        v_pos_candidates.extend([x1, x2])
                h_positions = _cluster_positions(h_pos_candidates, tol=6) if h_pos_candidates else []
                v_positions = _cluster_positions(v_pos_candidates, tol=6) if v_pos_candidates else []
                h_positions = sorted(list(set(h_positions)))
                v_positions = sorted(list(set(v_positions)))
                # 筛掉过于密集的噪声线（间距过小）
                h_positions = [p for i,p in enumerate(h_positions) if i==0 or (p - h_positions[i-1]) > 8]
                v_positions = [p for i,p in enumerate(v_positions) if i==0 or (p - v_positions[i-1]) > 8]
                if len(h_positions) >= 2 and len(v_positions) >= 2:
                    # 以线位置定义网格，逐格OCR
                    rows_text = []
                    # 确定表格区域边界
                    min_x, max_x = v_positions[0], v_positions[-1]
                    min_y, max_y = h_positions[0], h_positions[-1]
                    # 限制到图像范围
                    min_x = max(0, min_x); max_x = min(image.shape[1]-1, max_x)
                    min_y = max(0, min_y); max_y = min(image.shape[0]-1, max_y)
                    # 遍历相邻线形成的网格单元
                    for r in range(len(h_positions)-1):
                        y_top = h_positions[r]
                        y_bottom = h_positions[r+1]
                        row_cells = []
                        for c in range(len(v_positions)-1):
                            x_left = v_positions[c]
                            x_right = v_positions[c+1]
                            # 裁剪单元格区域，适当内缩，避免线条干扰
                            pad_x = max(1, (x_right - x_left) // 20)
                            pad_y = max(1, (y_bottom - y_top) // 20)
                            x1c = max(min_x, x_left + pad_x)
                            x2c = min(max_x, x_right - pad_x)
                            y1c = max(min_y, y_top + pad_y)
                            y2c = min(max_y, y_bottom - pad_y)
                            if x2c > x1c and y2c > y1c:
                                cell = image[y1c:y2c, x1c:x2c]
                                # 灰度与轻度二值化
                                cell_gray = cv2.cvtColor(cell, cv2.COLOR_BGR2GRAY) if cell.ndim == 3 else cell
                                cell_thr = cv2.adaptiveThreshold(cell_gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY, 11, 2)
                                # 单元格识别：单行/短文本
                                try:
                                    cell_text = pytesseract.image_to_string(cell_thr, lang='chi_sim+eng', config='--oem 3 --psm 7')
                                except Exception:
                                    cell_text = ''
                                row_cells.append(cell_text.strip())
                            else:
                                row_cells.append('')
                        rows_text.append(row_cells)
                    # 生成HTML与文本
                    table_text_structured = '\n'.join(['\t'.join(row) for row in rows_text]).strip()
                    table_html = self._generate_table_html([[None,None,None,None,cell] for row in rows_text for cell in row])
                    # 上面生成HTML需要每行结构，改为按行迭代
                    table_html = '<table border="1" cellpadding="4" cellspacing="0" style="border-collapse: collapse; width: 100%;">' + ''.join(['<tr>' + ''.join([f'<td>{cell}</td>' for cell in row]) + '</tr>' for row in rows_text]) + '</table>'
                    return {
                        'table_text': table_text_structured,
                        'table_html': table_html,
                        'has_table': True
                    }
            except Exception as e:
                logger.warning(f"网格表格提取失败，回退：{e}")
            
            # 创建一个表格线图像
            lines_image = np.zeros_like(gray)
            
            # 绘制水平线
            if horizontal_lines is not None:
                for line in horizontal_lines:
                    x1, y1, x2, y2 = line[0]
                    cv2.line(lines_image, (x1, y1), (x2, y2), 255, 2)
            
            # 绘制垂直线
            if vertical_lines is not None:
                for line in vertical_lines:
                    x1, y1, x2, y2 = line[0]
                    cv2.line(lines_image, (x1, y1), (x2, y2), 255, 2)
            
            # 合并所有线条
            table_mask = cv2.bitwise_or(detect_horizontal, detect_vertical)
            
            # 寻找轮廓
            contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            # 如果找到轮廓，尝试识别表格单元格
            if contours:
                # 选择最大的轮廓作为表格
                largest_contour = max(contours, key=cv2.contourArea)
                x, y, w, h = cv2.boundingRect(largest_contour)
                
                # 裁剪表格区域
                table_region = gray[y:y+h, x:x+w]
                
                # 使用Tesseract的表格识别配置
                table_config = '--oem 3 --psm 6'
                table_text = pytesseract.image_to_string(table_region, lang='chi_sim+eng', config=table_config)
                
                # 使用image_to_data获取更详细的文本位置信息
                data = pytesseract.image_to_data(table_region, output_type=pytesseract.Output.DICT)
                
                # 提取文本和位置信息
                text_with_positions = []
                for i in range(len(data['text'])):
                    if int(data['conf'][i]) > 0 and data['text'][i].strip():
                        x_pos = data['left'][i]
                        y_pos = data['top'][i]
                        width = data['width'][i]
                        height = data['height'][i]
                        text = data['text'][i]
                        text_with_positions.append((x_pos, y_pos, width, height, text))
                
                # 根据位置信息尝试重建表格结构
                # 简单的方法：按行和列排序
                if text_with_positions:
                    # 按y坐标排序（行）
                    text_with_positions.sort(key=lambda t: t[1])
                    
                    # 估算行高
                    avg_height = sum(t[3] for t in text_with_positions) / len(text_with_positions)
                    
                    # 分组行
                    rows = []
                    current_row = [text_with_positions[0]]
                    
                    for i in range(1, len(text_with_positions)):
                        if text_with_positions[i][1] - current_row[-1][1] < avg_height * 1.5:
                            current_row.append(text_with_positions[i])
                        else:
                            # 对当前行按x坐标排序
                            current_row.sort(key=lambda t: t[0])
                            rows.append(current_row)
                            current_row = [text_with_positions[i]]
                    
                    # 添加最后一行
                    if current_row:
                        current_row.sort(key=lambda t: t[0])
                        rows.append(current_row)
                    
                    # 生成表格文本（使用制表符分隔）
                    table_text_structured = ''
                    for row in rows:
                        row_text = '\t'.join(t[4] for t in row)
                        table_text_structured += row_text + '\n'
                    
                    # 生成HTML表格
                    table_html = self._generate_table_html(rows)
                    
                    return {
                        'table_text': table_text_structured.strip(),
                        'table_html': table_html,
                        'has_table': True
                    }
                else:
                    # 如果无法获取位置信息，返回原始识别文本
                    return {
                        'table_text': table_text.strip(),
                        'table_html': None,
                        'has_table': True
                    }
            else:
                # 尝试使用不同的配置进行表格识别
                try:
                    # 使用页面分割模式 4（假设是单个列文本）
                    table_text = pytesseract.image_to_string(
                        image, 
                        lang='chi_sim+eng', 
                        config='--oem 3 --psm 4'
                    )
                    
                    # 检查是否可能是表格
                    lines = table_text.strip().split('\n')
                    if len(lines) >= 2 and any(len(line.split()) >= 3 for line in lines):
                        return {
                            'table_text': table_text.strip(),
                            'table_html': None,
                            'has_table': True
                        }
                except Exception as e:
                    logger.warning(f"备用表格识别失败: {e}")
                
                # 如果无法识别表格，回退为纯OCR文本（保证总能有结果）
                try:
                    fallback_img = image
                    if crop_area and isinstance(crop_area, dict):
                        try:
                            fallback_img = self.crop_image(image, crop_area)
                        except Exception:
                            fallback_img = image
                    # 统一轻度预处理，提高可读性
                    proc = self.preprocess_for_ocr(fallback_img)
                    fallback_text = pytesseract.image_to_string(proc, lang='chi_sim+eng', config='--oem 3 --psm 6').strip()
                    logger.info(f"表格结构缺失，已回退为纯文本，长度: {len(fallback_text)}")
                    return {
                        'table_text': fallback_text if fallback_text else "",
                        'table_html': None,
                        'has_table': False
                    }
                except Exception as fe:
                    logger.warning(f"纯文本回退失败: {fe}")
                    return {
                        'table_text': "",
                        'table_html': None,
                        'has_table': False
                    }
        except Exception as e:
            logger.error(f"表格识别错误: {str(e)}")
            raise
    
    def _generate_table_html(self, rows):
        """根据识别的行生成HTML表格"""
        html = '<table border="1" cellpadding="4" cellspacing="0" style="border-collapse: collapse; width: 100%;">\n'
        
        for row in rows:
            html += '  <tr>\n'
            for cell in row:
                html += f'    <td>{cell[4]}</td>\n'
            html += '  </tr>\n'
        
        html += '</table>'
        return html
    
    def decode_image(self, base64_string):
        """解码base64图像数据"""
        try:
            # 确保输入是字符串
            if not isinstance(base64_string, str):
                raise ValueError("图像数据必须是字符串格式")
            
            # 移除前缀（如果有）
            if base64_string.startswith('data:image/'):
                base64_string = base64_string.split(',')[1]
            
            # 添加填充以确保base64解码正确
            padding = len(base64_string) % 4
            if padding:
                base64_string += '=' * (4 - padding)
            
            # 解码并转换为numpy数组
            img_data = base64.b64decode(base64_string)
            img = Image.open(BytesIO(img_data))
            
            # 转换为RGB格式
            if img.mode == 'RGBA':
                img = img.convert('RGB')
            
            return np.array(img)
        except Exception as e:
            logger.error(f"图像解码失败: {str(e)}")
            raise
    
    def rotate_image(self, img, angle, enhance=False):
        """旋转图像，并可选地增强图像"""
        if img is None:
            raise ValueError("输入图像为空")
        
        # 如果需要增强图像
        if enhance:
            # 转换为灰度图（如果是彩色图像）
            if len(img.shape) == 3:
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            else:
                gray = img
            
            # 使用高斯滤波降噪
            denoised = cv2.GaussianBlur(gray, (3, 3), 0)
            
            # 使用自适应直方图均衡化增强对比度
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            enhanced = clahe.apply(denoised)
            
            # 如果原图像是彩色的，转换回彩色
            if len(img.shape) == 3:
                img = cv2.cvtColor(enhanced, cv2.COLOR_GRAY2BGR)
            else:
                img = enhanced
        
        # 转换角度为浮点数
        angle = float(angle)
        
        h, w = img.shape[:2]
        center = (w // 2, h // 2)
        
        # 高效旋转90度倍数
        if abs(angle % 90) < 0.1:
            angle_rounded = round(angle / 90) * 90
            if angle_rounded == 90 or angle_rounded == -270:
                rotated = cv2.rotate(img, cv2.ROTATE_90_CLOCKWISE)
            elif angle_rounded == 180 or angle_rounded == -180:
                rotated = cv2.rotate(img, cv2.ROTATE_180)
            elif angle_rounded == 270 or angle_rounded == -90:
                rotated = cv2.rotate(img, cv2.ROTATE_90_COUNTERCLOCKWISE)
            else:
                rotated = img.copy()
        else:
            # 一般角度旋转
            M = cv2.getRotationMatrix2D(center, angle, 1.0)
            cos = np.abs(M[0, 0])
            sin = np.abs(M[0, 1])
            new_w = int((h * sin) + (w * cos))
            new_h = int((h * cos) + (w * sin))
            M[0, 2] += (new_w / 2) - center[0]
            M[1, 2] += (new_h / 2) - center[1]
            flags = cv2.INTER_LINEAR if abs(angle) > 10 else cv2.INTER_NEAREST
            rotated = cv2.warpAffine(
                img, M, (new_w, new_h),
                flags=flags,
                borderMode=cv2.BORDER_CONSTANT,
                borderValue=255
            )
        
        logger.info(f"图像旋转完成: {angle}度")
        return rotated
    
    def preprocess_for_ocr(self, img):
        """自适应图像预处理：放大、去噪、对比度增强、自动纠偏、阈值与形态学，提升OCR准确率"""
        try:
            # 1) 转为灰度
            if img.ndim == 3:
                gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
            else:
                gray = img.copy()

            h, w = gray.shape[:2]
            # 2) 小图自动放大（最多放大到较长边1600像素）
            max_side = max(h, w)
            if max_side < 800:
                scale = min(1600 / max_side, 2.0)
                new_w = int(w * scale)
                new_h = int(h * scale)
                gray = cv2.resize(gray, (new_w, new_h), interpolation=cv2.INTER_CUBIC)
                h, w = new_h, new_w

            # 3) 轻度去噪 + 局部对比度增强
            denoised = cv2.bilateralFilter(gray, d=5, sigmaColor=50, sigmaSpace=50)
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            enhanced = clahe.apply(denoised)

            # 4) 基于最小外接矩形估计倾斜角并纠偏（避免过度旋转）
            try:
                thr = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
                inv = cv2.bitwise_not(thr)
                coords = np.column_stack(np.where(inv > 0))
                if coords.shape[0] > 500:  # 足够的前景像素
                    rect = cv2.minAreaRect(coords.astype(np.float32))
                    angle = rect[-1]
                    if angle < -45:
                        angle = -(90 + angle)
                    else:
                        angle = -angle
                    if abs(angle) >= 0.5 and abs(angle) <= 8.0:
                        M = cv2.getRotationMatrix2D((w / 2, h / 2), angle, 1.0)
                        enhanced = cv2.warpAffine(enhanced, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
            except Exception:
                pass

            # 5) 根据直方图选择阈值方法：Otsu优先，极端情况回退自适应
            try:
                _, otsu_bin = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                black_ratio = (otsu_bin == 0).sum() / (otsu_bin.size)
                if 0.05 < black_ratio < 0.95:
                    binary = otsu_bin
                else:
                    binary = cv2.adaptiveThreshold(
                        enhanced, 255,
                        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                        cv2.THRESH_BINARY,
                        11, 2
                    )
            except Exception:
                binary = cv2.adaptiveThreshold(
                    enhanced, 255,
                    cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                    cv2.THRESH_BINARY,
                    11, 2
                )

            # 6) 形态学处理：先开后闭，去小噪点并连接细字符
            kernel_small = np.ones((1, 1), np.uint8)
            kernel_mid = np.ones((2, 2), np.uint8)
            opened = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel_small)
            processed = cv2.morphologyEx(opened, cv2.MORPH_CLOSE, kernel_mid)

            # 7) 轻度模糊，抑制锯齿
            processed = cv2.GaussianBlur(processed, (1, 1), 0)
            return processed
        except Exception as e:
            logger.error(f"图像预处理失败: {str(e)}")
            # 失败时返回原始灰度图作为备用
            if img.ndim == 3:
                return cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
            return img

    def enhance_image(self, img):
        """增强版图像增强，大幅提高对比度和清晰度"""
        try:
            # 先使用OpenCV进行基础增强
            # 转换为灰度图进行处理
            if len(img.shape) == 3:
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            else:
                gray = img.copy()
            
            # 1. 自适应直方图均衡化 - 增强局部对比度
            clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
            clahe_enhanced = clahe.apply(gray)
            
            # 2. 非局部均值去噪 - 减少噪点但保留边缘
            denoised = cv2.fastNlMeansDenoising(clahe_enhanced, h=10)
            
            # 3. 锐化处理 - 提高清晰度
            kernel = np.array([[-1, -1, -1], [-1, 9, -1], [-1, -1, -1]])
            sharpened = cv2.filter2D(denoised, -1, kernel)
            
            # 4. 使用PIL进行更高级的增强
            if len(img.shape) == 3:
                # 如果原图是彩色的，先转换回彩色
                enhanced_bgr = cv2.cvtColor(sharpened, cv2.COLOR_GRAY2BGR)
                pil_img = Image.fromarray(enhanced_bgr)
            else:
                pil_img = Image.fromarray(sharpened)
            
            # 5. 增强对比度 - 更强的对比度调整
            enhancer = ImageEnhance.Contrast(pil_img)
            img_contrast = enhancer.enhance(1.8)
            
            # 6. 增强亮度
            enhancer = ImageEnhance.Brightness(img_contrast)
            img_bright = enhancer.enhance(1.3)
            
            # 7. 增强锐度
            enhancer = ImageEnhance.Sharpness(img_bright)
            img_sharp = enhancer.enhance(2.0)
            
            enhanced_img = np.array(img_sharp)
            logger.info(f"图像增强完成，增强后尺寸: {enhanced_img.shape}")
            return enhanced_img
        except Exception as e:
            logger.error(f"图像增强失败: {str(e)}")
            return img  # 失败时返回原图
    
    def ocr_text_merged(self, img, crop_area=None, top_k: int = 3):
        """全新OCR实现：单一入口，内部自包含多路预处理与多配置尝试，
        返回最佳结果与若干候选文本。
        返回: { 'text': str, 'alternatives': [str, ...] }
        """
        if img is None:
            raise ValueError("输入图像为空")

        try:
            # 1) 可选裁剪（前端通常已裁剪，但仍支持后端裁剪）
            if crop_area:
                img = self.crop_image(img, crop_area)

            # 2) 轻量方向检测与校正（尽量容错）
            try:
                h0, w0 = img.shape[:2]
                if max(h0, w0) > 600:
                    gray0 = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY) if img.ndim == 3 else img
                    osd = pytesseract.image_to_osd(Image.fromarray(gray0))
                    m = re.search(r"Rotate:\s*(\d+)", osd)
                    if m:
                        deg = int(m.group(1)) % 360
                        if deg in (90, 180, 270):
                            img = self.rotate_image(img, deg, enhance=False)
                            logger.info(f"自动方向校正: 旋转 {deg} 度")
            except Exception as e:
                logger.warning(f"方向检测失败或跳过: {e}")

            # 3) 构造预处理分支（纯灰度、自适应阈值、OTSU、反锐化 + 归一化）
            if img.ndim == 3:
                base_gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            else:
                base_gray = img.copy()

            def unsharp(gray: np.ndarray) -> np.ndarray:
                blur = cv2.GaussianBlur(gray, (0, 0), 1.0)
                sharp = cv2.addWeighted(gray, 1.5, blur, -0.5, 0)
                return sharp

            def norm_contrast(gray: np.ndarray) -> np.ndarray:
                return cv2.normalize(gray, None, 0, 255, cv2.NORM_MINMAX)

            # 自适应阈值与OTSU
            try:
                adap = cv2.adaptiveThreshold(base_gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                             cv2.THRESH_BINARY, 31, 10)
            except Exception:
                adap = base_gray
            try:
                _, otsu = cv2.threshold(base_gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            except Exception:
                otsu = base_gray

            variants = []
            variants.append(("gray", base_gray))
            variants.append(("adap", adap))
            variants.append(("otsu", otsu))
            variants.append(("sharp_norm", norm_contrast(unsharp(base_gray))))

            # 小图放大，边缘加白边
            tuned_variants = []
            for name, im in variants:
                h, w = im.shape[:2]
                scale = 1.0
                if max(h, w) < 900:
                    scale = min(2.5, 900.0 / max(h, w))
                if scale > 1.05:
                    im = cv2.resize(im, (int(w * scale), int(h * scale)), interpolation=cv2.INTER_CUBIC)
                try:
                    im = cv2.copyMakeBorder(im, 8, 12, 8, 8, cv2.BORDER_CONSTANT, value=255)
                except Exception:
                    pass
                tuned_variants.append((name, im))

            # 4) 语言初判（快速小样本）：估计中文占比，决定语言优先顺序与空格策略
            def count_chars(s: str):
                if not s:
                    return 0, 0
                cjk = len(re.findall(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]", s))
                latin = len(re.findall(r"[A-Za-z]", s))
                return cjk, latin

            try:
                probe = cv2.resize(base_gray, (min(600, base_gray.shape[1]), int(base_gray.shape[0] * min(600, base_gray.shape[1]) / max(1, base_gray.shape[1]))))
            except Exception:
                probe = base_gray
            probe_img = Image.fromarray(probe)

            probe_txt_zh = ""
            probe_txt_en = ""
            try:
                probe_txt_zh = pytesseract.image_to_string(probe_img, lang='chi_sim', config='--oem 3 --psm 7 -c user_defined_dpi=300')
            except Exception:
                pass
            try:
                probe_txt_en = pytesseract.image_to_string(probe_img, lang='eng', config='--oem 3 --psm 7 -c user_defined_dpi=300 -c preserve_interword_spaces=1')
            except Exception:
                pass
            cjk_cnt, latin_cnt = count_chars((probe_txt_zh or '') + (probe_txt_en or ''))
            total_alpha = max(1, cjk_cnt + latin_cnt)
            cjk_ratio = cjk_cnt / total_alpha

            if cjk_ratio >= 0.3:
                lang_order = ['chi_sim+eng', 'chi_sim', 'eng']
            else:
                lang_order = ['eng', 'chi_sim+eng', 'chi_sim']

            # 5) 多配置尝试（全新评分体系）：综合置信度、行数、长度与语言偏好
            psms = [6, 3, 4, 7, 11]
            results_meta = []

            for vname, im in tuned_variants:
                pil_img = Image.fromarray(im)
                for lg in lang_order:
                    preserve_spaces = 0 if ('chi_sim' in lg and cjk_ratio >= 0.2) else 1
                    for psm in psms:
                        cfg = f"--oem 3 --psm {psm} -c user_defined_dpi=300 -c preserve_interword_spaces={preserve_spaces}"
                        try:
                            txt = pytesseract.image_to_string(pil_img, lang=lg, config=cfg) or ""
                            data = pytesseract.image_to_data(pil_img, lang=lg, config=cfg, output_type=Output.DICT)
                            conf_vals = []
                            for c in data.get('conf', []) or []:
                                try:
                                    ci = int(c)
                                    if ci >= 0:
                                        conf_vals.append(ci)
                                except Exception:
                                    continue
                            mean_conf = float(sum(conf_vals)) / len(conf_vals) if conf_vals else 0.0
                            word_count = len([t for t in (data.get('text', []) or []) if t and t.strip()])
                            line_count = max(data.get('line_num', [0] or [0])) if isinstance(data.get('line_num'), list) else 0
                            cjk_c, latin_c = count_chars(txt)

                            # 评分：置信度 + 语言偏好 + 行数/长度适度加分
                            lang_boost = 0.0
                            if cjk_ratio >= 0.3:
                                lang_boost += 6.0 if 'chi_sim' in lg else (-4.0 if lg == 'eng' else 0.0)
                            else:
                                lang_boost += 3.0 if lg == 'eng' else 0.5 if 'chi_sim+eng' in lg else 0.0

                            length_boost = min(5.0, len(txt) / 200.0)
                            line_boost = min(3.0, line_count / 10.0)

                            score = round(mean_conf + lang_boost + length_boost + line_boost, 3)
                            results_meta.append({
                                'text': txt,
                                'variant': vname,
                                'lang': lg,
                                'psm': psm,
                                'score': score,
                                'mean_conf': mean_conf,
                                'word_count': word_count,
                                'cjk': cjk_c,
                                'latin': latin_c,
                            })
                        except Exception as e:
                            logger.warning(f"OCR失败 variant={vname} lang={lg} psm={psm}: {e}")

            if not results_meta:
                # 彻底兜底
                try:
                    txt = pytesseract.image_to_string(Image.fromarray(base_gray)) or ""
                    results_meta.append({'text': txt, 'variant': 'fallback', 'lang': 'auto', 'psm': -1, 'score': 0.0, 'mean_conf': 0.0, 'word_count': len(txt.split()), 'cjk': 0, 'latin': 0})
                except Exception as e:
                    logger.error(f"全部OCR尝试失败: {e}")
                    raise

            # 6) 文本清理和去重
            def clean_text_new(s: str) -> str:
                if not s:
                    return ''
                s = s.replace('\r\n', '\n').replace('\r', '\n')
                s = re.sub(r'[\t\x0b\x0c]', ' ', s)
                s = re.sub(r'-\n', '', s)  # 合并断字
                s = re.sub(r'\s+\n', '\n', s)
                s = re.sub(r'\n{3,}', '\n\n', s)
                s = re.sub(r'\s{2,}', ' ', s)
                return s.strip()

            for r in results_meta:
                r['text'] = clean_text_new(r['text'])

            # 唯一化（严格去重）
            uniq = []
            for r in sorted(results_meta, key=lambda x: x['score'], reverse=True):
                t = r['text']
                if not t:
                    continue
                if all(t != u['text'] for u in uniq):
                    uniq.append(r)
                if len(uniq) >= top_k + 1:
                    break

            best = uniq[0]['text'] if uniq else ''
            alts = [u['text'] for u in uniq[1:]]

            logger.info(f"新OCR完成：主结果len={len(best)} 备选数={len(alts)} cjk_ratio={cjk_ratio:.2f}")
            return {'text': best, 'alternatives': alts}
        except Exception as e:
            logger.error(f"OCR识别失败: {str(e)}")
            raise

    def fix_text(self, text):
        """增强版文本纠错"""
        if not text:
            return ""
        
        # 扩展的纠错规则
        fix_rules = {
            # 中文常见错误
            "分折": "分析", "工贝": "工具", "即然": "既然", "象": "像", "做": "作",
            "仃": "停", "亍": "行", "彳": "行", "攵": "文", "肀": "肀",
            "匚": "区", "冂": "同", "凵": "山", "爫": "爪", "丬": "壮",
            # 英文常见错误
            "teh": "the", "hwo": "how", "tahn": "than", "whta": "what",
            "adn": "and", "tihng": "thing", "waht": "what", "sihde": "side",
            # 标点符号转换
            "，": ",", "。": ".", "；": ";", "：": ":", "？": "?",
            "！": "!", "（": "(", "）": ")", "【": "[", "】": "]",
            # 常见空格问题
            "\t": " ", "  ": " "
        }
        
        fixed = text
        for wrong, right in fix_rules.items():
            fixed = fixed.replace(wrong, right)
        
        # 清理多余的空行
        fixed = re.sub(r'\n{3,}', '\n\n', fixed)
        
        logger.info(f"文本纠错完成")
        return fixed
    
    def encode_image(self, img):
        """将图像编码为base64"""
        if img is None:
            raise ValueError("输入图像为空")
        
        try:
            # 确保图像格式正确
            if img.dtype != np.uint8:
                img = img.astype(np.uint8)
            
            # 转换为PIL图像
            if len(img.shape) == 3:
                pil_img = Image.fromarray(img)
            else:
                pil_img = Image.fromarray(img, mode='L')
            
            # 保存到内存
            buffer = BytesIO()
            pil_img.save(buffer, format="JPEG", quality=90, optimize=True)
            buffer.seek(0)
            
            # 编码为base64
            encoded = base64.b64encode(buffer.read()).decode('utf-8')
            logger.info("图像编码完成")
            return encoded
        except Exception as e:
            logger.error(f"图像编码失败: {str(e)}")
            raise

# 全局识别器实例
recognizer = DocRecognizer()

# LLM客户端（如豆包）初始化
try:
    from llm_client import LLMClient
    llm_client = LLMClient.from_env()
    logger.info('LLM客户端初始化完成')
except Exception as e:
    llm_client = None
    logger.warning(f'LLM客户端初始化失败或未配置: {str(e)}')

# 远程OCR客户端
try:
    from remote_ocr_client import RemoteOCRClient
    remote_ocr = RemoteOCRClient.from_env()
    logger.info('远程OCR客户端初始化完成')
except Exception as e:
    remote_ocr = None
    logger.warning(f'远程OCR未配置或初始化失败: {str(e)}')

# 健康检查端点
@app.route('/health', methods=['GET'])
def health_check():
    """健康检查API"""
    return jsonify({
        'status': 'healthy',
        'service': 'DocVision AI',
        'version': '1.0.0'
    }), 200

# 静态文件服务
@app.route('/')
def index():
    """提供前端页面"""
    try:
        return send_from_directory(app.static_folder, 'index.html')
    except Exception as e:
        logger.error(f"提供前端页面失败: {str(e)}")
        return jsonify({'error': '无法提供前端页面'}), 500

# 旋转图像API
@app.route('/api/rotate', methods=['POST'])
def rotate():
    """旋转图像API"""
    try:
        data = request.json
        if not data or 'image' not in data or 'angle' not in data:
            return jsonify({'error': '缺少必要参数: image 和 angle'}), 400
        
        logger.info("收到图像旋转请求")
        
        # 解码图像
        img = recognizer.decode_image(data['image'])
        
        # 旋转图像，支持可选的增强功能
        enhance = data.get('enhance', False)
        rotated = recognizer.rotate_image(img, data['angle'], enhance)
        
        # 编码回base64
        encoded = recognizer.encode_image(rotated)
        
        return jsonify({
            'success': True,
            'image': f'data:image/jpeg;base64,{encoded}',
            'message': '图像旋转成功'
        })
    except ValueError as e:
        logger.warning(f"旋转参数错误: {str(e)}")
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        logger.error(f"图像旋转失败: {str(e)}")
        return jsonify({'error': f'图像旋转失败: {str(e)}'}), 500

# OCR识别API
@app.route('/api/ocr', methods=['POST'])
def ocr():
    """OCR识别API（支持variant=baseline/enhanced参数用于对比）"""
    try:
        data = request.json
        if not data or 'image' not in data:
            return jsonify({'error': '缺少必要参数: image'}), 400
        
        logger.info("收到OCR识别请求")
        
        # 解码图像
        img = recognizer.decode_image(data['image'])
        logger.info(f"图像解码成功，尺寸: {img.shape}")
        
        # 获取裁剪区域参数
        crop_area = data.get('crop_area', None)
        logger.info(f"裁剪区域参数: {crop_area}")
        
        # 执行OCR
        engine = (data.get('engine') or '').lower()
        variant = (data.get('variant') or '').lower()
        if engine in ('doubao', 'remote'):
            if remote_ocr is None:
                return jsonify({'error': '远程OCR未配置'}), 400
            # 注意：前端发送的是裁剪后的PNG base64（不含头部），我们直接透传
            text = remote_ocr.ocr(data['image'], crop_area=crop_area)
        else:
            # 统一本地识别：忽略 variant，使用合并后的单一管线，并返回备选
            merged = recognizer.ocr_text_merged(img, crop_area=crop_area)
            text = merged.get('text', '')
            alternatives = merged.get('alternatives', [])
        
        if not text:
            logger.warning("OCR识别未返回任何文本")
            return jsonify({
                'success': True,
                'text': '',
                'alternatives': alternatives if 'alternatives' in locals() else [],
                'message': '未识别到文本'
            })
        
        # 可选的文本纠错（对主结果应用）
        if data.get('auto_fix', False):
            text = recognizer.fix_text(text)
        
        return jsonify({
            'success': True,
            'text': text,
            'alternatives': alternatives if 'alternatives' in locals() else [],
            'message': 'OCR识别成功',
            'char_count': len(text)
        })
    except ValueError as e:
        logger.warning(f"OCR参数错误: {str(e)}")
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        logger.error(f"OCR识别失败: {str(e)}")
        return jsonify({'error': f'OCR识别失败: {str(e)}'}), 500

# 表格识别API
@app.route('/api/recognize-table', methods=['POST'])
def recognize_table():
    """表格识别API"""
    try:
        data = request.json
        if not data or 'image' not in data:
            return jsonify({'error': '缺少必要参数: image'}), 400
        
        logger.info("收到表格识别请求")
        
        # 解码图像
        img = recognizer.decode_image(data['image'])
        logger.info(f"图像解码成功，尺寸: {img.shape}")
        
        # 获取裁剪区域参数
        crop_area = data.get('crop_area', None)
        logger.info(f"裁剪区域参数: {crop_area}")
        
        # 执行表格识别
        table_result = recognizer.recognize_table(img, crop_area=crop_area)
        
        return jsonify({
            'success': True,
            'table_text': table_result['table_text'],
            'table_html': table_result.get('table_html'),
            'has_table': table_result['has_table'],
            'message': '表格识别成功'
        })
    except ValueError as e:
        logger.warning(f"表格识别参数错误: {str(e)}")
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        logger.error(f"表格识别失败: {str(e)}")
        return jsonify({'error': f'表格识别失败: {str(e)}'}), 500

# 文本纠错API
@app.route('/api/fix-text', methods=['POST'])
def fix_text():
    """文本纠错API（支持use_llm开关）"""
    try:
        data = request.json
        if not data or 'text' not in data:
            return jsonify({'error': '缺少文本参数'}), 400
        text = data['text']
        if not isinstance(text, str):
            return jsonify({'error': '文本必须是字符串格式'}), 400

        use_llm = bool(data.get('use_llm')) and llm_client is not None
        logger.info(f"收到文本纠错请求，use_llm={use_llm}")

        if use_llm:
            extra = data.get('instructions')
            fixed_text = llm_client.fix_text(text, extra_instructions=extra)
        else:
            fixed_text = recognizer.fix_text(text)

        return jsonify({
            'success': True,
            'text': fixed_text,
            'message': '文本纠错成功' + ("（LLM）" if use_llm else "（规则）"),
            'original_length': len(text),
            'fixed_length': len(fixed_text)
        })
    except Exception as e:
        logger.error(f"文本纠错失败: {str(e)}")
        return jsonify({'error': f'文本纠错失败: {str(e)}'}), 500

# LLM结构化抽取API
@app.route('/api/llm/extract', methods=['POST'])
def llm_extract():
    """使用LLM对OCR文本进行结构化抽取，返回JSON。"""
    if llm_client is None:
        return jsonify({'error': 'LLM未配置'}), 400
    try:
        data = request.json or {}
        text = data.get('text', '')
        schema = data.get('schema')
        instruction = data.get('instruction')
        if not text:
            return jsonify({'error': '缺少文本参数'}), 400
        result = llm_client.extract_structured(text, schema_hint=schema, instruction=instruction)
        return jsonify({'success': True, 'data': result})
    except Exception as e:
        return jsonify({'error': f'抽取失败: {str(e)}'}), 500

# 保存为Word文档
@app.route('/api/save-word', methods=['POST'])
def api_save_word():
    """保存OCR结果为Word文档"""
    try:
        data = request.json
        text = data.get('text', '')
        has_table = data.get('has_table', False)
        
        logger.info("收到保存Word文档请求")
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        doc.add_heading('OCR识别结果', level=1)
        
        # 添加正文内容
        # 按段落分割文本
        paragraphs = text.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
            else:
                # 保留空行
                doc.add_paragraph()
        
        # 如果有表格，可以在这里添加表格
        # 注意：由于前端仅发送has_table标志，实际表格数据需要额外处理
        if has_table:
            doc.add_heading('识别的表格', level=2)
            doc.add_paragraph('表格数据已在上方文本中显示')
        
        # 保存文档到内存
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # 返回文档
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f'ocr_result_{int(time.time())}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        logger.error(f"生成Word文档错误: {str(e)}")
        return jsonify({'error': f'生成Word文档失败: {str(e)}'}), 500

# 处理404错误
@app.errorhandler(404)
def not_found(error):
    return jsonify({
        'error': '路径不存在',
        'message': '请检查请求路径是否正确'
    }), 404

# 处理500错误
@app.errorhandler(500)
def internal_error(error):
    logger.error(f"内部服务器错误: {str(error)}")
    return jsonify({
        'error': '内部服务器错误',
        'message': '服务器遇到了一个错误，请稍后再试'
    }), 500

if __name__ == '__main__':
    # 应用配置
    app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0  # 禁用缓存
    
    # 确保前端目录存在
    if not os.path.exists(app.static_folder):
        os.makedirs(app.static_folder)
        logger.warning(f"创建前端目录: {app.static_folder}")
    
    logger.info("启动DocVision AI服务...")
    logger.info(f"服务地址: http://0.0.0.0:5000")
    logger.info(f"静态文件目录: {app.static_folder}")
    
    # 以调试模式启动服务
    app.run(host='0.0.0.0', port=5000, debug=True, threaded=True)